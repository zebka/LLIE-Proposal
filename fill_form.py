# -*- coding: utf-8 -*-
"""Fill the Birjand PhD proposal form (PhDProposalform.docx) with the
Twilight-LLIE proposal drafts. Produces PhDProposalform-Completed.docx"""
import re
import io
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(r"C:\Users\Alireza\Desktop\New folder (2)")
PROJ = BASE / "Twilight-LLIE-Proposal"
FORM = BASE / "PhDProposalform.docx"
OUT = BASE / "PhDProposalform-Completed.docx"

# ---------------------------------------------------------------- bib parsing
def parse_bib(path):
    text = path.read_text(encoding="utf-8")
    entries = {}
    for m in re.finditer(r"@(\w+)\{([^,\s]+),\s*\n(.*?)\n\}", text, re.S):
        kind, key, body = m.group(1), m.group(2), m.group(3)
        fields = {}
        for fm in re.finditer(r"(\w+)\s*=\s*\{(.+?)\}\s*,?\s*\n", body + "\n", re.S):
            fields[fm.group(1).lower()] = re.sub(r"\s+", " ", fm.group(2)).strip()
        fields["kind"] = kind
        entries[key] = fields
    return entries


def fmt_authors(author_field):
    names = [n.strip() for n in author_field.split(" and ")]
    formatted = []
    for n in names:
        if "," in n:
            last, first = [p.strip() for p in n.split(",", 1)]
            initials = ". ".join(w[0] for w in first.split() if w) + "."
            formatted.append(f"{initials} {last}")
        else:
            parts = n.split()
            if len(parts) == 1:
                formatted.append(parts[0])
            else:
                initials = ". ".join(w[0] for w in parts[:-1]) + "."
                formatted.append(f"{initials} {parts[-1]}")
    if len(formatted) > 6:
        return formatted[0] + " et al."
    if len(formatted) > 1:
        return ", ".join(formatted[:-1]) + ", and " + formatted[-1]
    return formatted[0]


def ieee_ref(key, entries):
    e = entries[key]
    authors = fmt_authors(e.get("author", ""))
    title = re.sub(r"[{}]", "", e.get("title", "")).strip()
    year = e.get("year", "")
    if e["kind"] == "inproceedings":
        venue = e.get("booktitle", "")
        s = f'{authors}, "{title}," in {venue}, {year}.'
    else:
        venue = e.get("journal", "")
        s = f'{authors}, "{title}," {venue}'
        if "volume" in e:
            s += f', vol. {e["volume"]}'
        if "number" in e:
            s += f', no. {e["number"]}'
        if "pages" in e:
            s += f', pp. {e["pages"].replace("--", "-")}'
        s += f", {year}."
    return s


# ------------------------------------------------------------ md cleaning
CITE_RE = re.compile(r"\[([^\[\]]+)\]")
KEY_RE = re.compile(r"^[A-Za-z][A-Za-z0-9_]*$")


def clean_md(md_text):
    """md -> list of (kind, text) where kind in {h, p, sp}"""
    items = []
    for raw in md_text.splitlines():
        s = raw.strip()
        if not s:
            if items and items[-1][0] != "sp":
                items.append(("sp", ""))
            continue
        if s.startswith(">") or s == "---" or s.startswith("%"):
            continue
        if s.startswith("[[TABLE") and s.endswith("]]"):
            items.append(("table", s[2:-2].replace("TABLE", "").strip()))
            continue
        if s.startswith("|"):
            continue  # md tables are hardcoded separately
        s = s.replace("**", "").replace("`", "").replace("*", "")
        if s.startswith("#### "):
            s = s[5:]
            items.append(("h", s))
        elif s.startswith("### "):
            s = s[4:]
            items.append(("h", s))
        elif s.startswith("## "):
            s = s[3:]
            items.append(("h", s))
        elif s.startswith("# "):
            s = s[2:]
            items.append(("h", s))
        elif s.startswith("- "):
            items.append(("p", "\u2022 " + s[2:]))
        else:
            items.append(("p", s))
    return items


def build_cite_map(sections_items, entries):
    order = {}
    for items in sections_items:
        for _, txt in items:
            for m in CITE_RE.finditer(txt):
                for tok in re.split(r"[،,]", m.group(1)):
                    tok = tok.strip()
                    if KEY_RE.match(tok) and tok in entries and tok not in order:
                        order[tok] = len(order) + 1
    return order


def map_cites(txt, order):
    def repl(m):
        toks = []
        for tok in re.split(r"[،,]", m.group(1)):
            tok = tok.strip()
            if KEY_RE.match(tok) and tok in order:
                toks.append(str(order[tok]))
            elif tok:
                toks.append(tok)
        return "[" + ", ".join(toks) + "]"
    return CITE_RE.sub(repl, txt)


def apply_cites(items, order):
    return [(kind, map_cites(txt, order)) for kind, txt in items]


# ------------------------------------------------------------ docx helpers
def set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def add_para(cell, text, rtl=True, bold=False, size=10.5):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_ltr_para(cell, text, size=9):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def fill_cell(cell, items, size=10.5):
    first = True
    for kind, txt in items:
        if kind == "sp":
            sp = cell.add_paragraph()
            sp.paragraph_format.space_after = Pt(2)
            first = False
            continue
        if kind == "table":
            fill_table(cell, TABLES.get(txt, [["?", "?"]]))
            first = False
            continue
        if first and not cell.paragraphs[0].text.strip():
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_rtl(p)
            run = p.add_run(txt)
            run.bold = kind == "h"
            run.font.size = Pt(size)
            first = False
        else:
            add_para(cell, txt, bold=(kind == "h"), size=size)


def fill_table(cell, rows, header=True, size=9):
    t = cell.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        t.style = "Table Grid"
    except Exception:
        set_table_borders(t)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri == 0 and header:
                set_rtl(p)
            run = p.add_run(val)
            run.bold = ri == 0 and header
            run.font.size = Pt(size)
    return t


# ------------------------------------------------------------ hardcoded tables
RISK_ROWS = [
    ["ریسک", "راهکار"],
    ["تأخیر در جمع‌آوری دیتاست TwilightDrive", "روش و مقاله اول صرفاً با بنچمارک‌های موجود منتشر می‌شود؛ دیتاست در فاز ۳–۴ (بخش ۶) قرار دارد"],
    ["محدودیت حافظه GPU برای استنتاج diffusion", "استنتاج روی فضای نهفته/موجک [jiang2023diffll] + اجرای سنگین‌ها روی Colab L4"],
    ["عدم تحقق کامل فرضیه ۲", "تحلیل خرابی به‌عنوان یافته علمی مستقل (سوال فرعی س۱) گزارش می‌شود"],
]

COMPARISON_ROWS = [
    ["قابلیت", "Cho et al.", "He et al.", "Self-DACE++", "روش پیشنهادی"],
    ["بدون جفت‌داده", "بله", "بله", "بله", "بله"],
    ["بدون آموزش وزن", "بله", "بله", "خیر", "بله"],
    ["prior فرکانسی", "خیر", "بله", "خیر", "بله"],
    ["سازگاری با نوردهی نامتوازن", "خیر", "خیر", "خیر", "بله (گیت روشنایی)"],
    ["ارزیابی در رژیم گرگ‌ومیش", "خیر", "خیر", "خیر", "بله"],
    ["دیتاست اختصاصی رژیم واسط", "خیر", "خیر", "خیر", "بله (TwilightDrive)"],
    ["ارزیابی وظیفه پایین‌دستی در گرگ‌ومیش", "خیر", "خیر", "خیر", "بله"],
    ["استنتاج تطبیقی سبک‌وزن", "خیر", "خیر", "بله", "بله"],
]

TIMELINE_ROWS = [
    ["فاز", "ماه", "فعالیت", "خروجی مورد انتظار"],
    ["۰", "۱–۳", "تکمیل مرور ادبیات؛ پیاده‌سازی و بازتولید baselineها؛ ساخت مجموعه ارزیابی اولیه گرگ‌ومیش برای تحلیل خرابی", "پروپوزال نهایی، کد baseline، گزارش تحلیل خرابی (پاسخ س۱)"],
    ["۱", "۴–۸", "طراحی و پیاده‌سازی ماژول گیت روشنایی و prior فرکانسی؛ ارزیابی اولیه روی بنچمارک‌های موجود", "مدل ساده + مقاله کنفرانسی"],
    ["۲", "۹–۱۴", "توسعه کامل سه ماژول؛ ablation و بهینه‌سازی استنتاج تطبیقی؛ ارزیابی جامع روی بنچمارک‌های استاندارد", "نسخه میانی روش + پیش‌نویس مقاله مجله اول"],
    ["۳", "۱۵–۱۹", "ضبط و پالایش (کیوریشن) TwilightDrive؛ ارزیابی پایین‌دستی (تشخیص شیء)", "مقاله مجله اول + نسخه اولیه دیتاست"],
    ["۴", "۲۰–۲۴", "تکمیل، ناشناس‌سازی و در صورت احراز شرایط، انتشار دیتاست و کد؛ مقاله دوم؛ نگارش رساله", "مقاله دوم + دیتاست منتشرشده + رساله تکمیل‌شده"],
]

MILESTONE = ("نقاط کنترل: پایان فاز ۱ = تصمیم go/no-go بر اساس نتایج فرضیه ۱؛ "
             "پایان فاز ۳ = ارزیابی فرضیه‌های ۲ و ۳؛ پایان فاز ۴ = دفاع بر اساس دو مقاله + دیتاست عمومی.")

TABLES = {
    "1": [
        ["دوره", "رویکرد", "نمایندگان", "محدودیت اصلی"],
        ["1971–1977", "تئوری Retinex", "Land و McCann [land1977retinex]", "چارچوب نظری؛ حل تحلیلی به‌دلیل عدم یکتایی تجزیه ممکن نیست"],
        ["دهه 1980", "تراز هیستوگرام تطبیقی", "AHE/CLAHE [pizer1987clahe]", "تقویت نویز و اغتشاش رنگ در نواحی ناهموار"],
        ["1997", "Retinex چندمقیاس", "SSR/MSRCR [jobson1997msrcr]", "حساسیت به پارامترهای گوسی و رنگ مصنوعی"],
        ["2011", "قیاس با برداشتن غبار", "Dong et al. [dong2011dehaze]", "فرض مدل پراکندگی در صحنه کم‌نور واقعی برقرار نیست"],
        ["2016", "مدل واریانس‌پذیر وزن‌دار", "SRIE [fu2016srie]", "حل‌کننده تکراری کند؛ تنظیم دستی پارامترها"],
        ["2017", "تخمین نگاشت روشنایی", "LIME [guo2017lime]", "وابستگی به prior ساختاری دست‌ساز"],
        ["2017", "نخستین یادگیری عمیق", "LLNet [lore2017llnet]", "ظرفیت محدود autoencoder و آموزش ناپایدار"],
    ],
    "2": [
        ["روش", "LOL-v1 (PSNR/SSIM)", "LOL-v2-real (PSNR/SSIM)", "LOL-v2-syn (PSNR/SSIM)"],
        ["RetinexNet", "16.77/0.560", "15.47/0.567", "17.13/0.798"],
        ["DeepUPE", "14.38/0.446", "13.27/0.452", "15.08/0.623"],
        ["EnlightenGAN", "17.48/0.650", "18.23/0.617", "16.57/0.734"],
        ["RUAS", "18.23/0.720", "18.37/0.723", "16.55/0.652"],
        ["DRBN", "20.13/0.830", "20.29/0.831", "23.22/0.927"],
        ["KinD", "20.86/0.790", "14.74/0.641", "13.29/0.578"],
        ["Restormer", "22.43/0.823", "19.94/0.827", "21.41/0.830"],
        ["MIRNet", "24.14/0.830", "20.02/0.820", "21.94/0.876"],
        ["SNR-Net", "24.61/0.842", "21.48/0.849", "24.14/0.928"],
        ["Retinexformer", "25.16/0.845", "22.80/0.840", "25.67/0.930"],
    ],
    "3": [
        ["روش", "اندازه مدل (M)", "FLOPs (G)", "زمان استنتاج (s)"],
        ["RetinexNet", "0.84", "136.02", "0.119"],
        ["KinD", "8.54", "29.13", "0.181"],
        ["DRBN", "0.58", "37.79", "0.053"],
        ["EnlightenGAN", "8.64", "61.01", "0.0097"],
        ["Zero-DCE", "0.079", "5.21", "0.0042"],
        ["RUAS", "0.0014", "0.28", "0.0063"],
        ["SCI", "0.0003", "0.062", "0.0017"],
    ],
    "4": [
        ["روش", "LOL PSNR", "LOL SSIM", "LOL LPIPS", "v2-real PSNR", "v2-real SSIM"],
        ["Zero-DCE", "14.861", "0.562", "0.335", "18.059", "0.580"],
        ["KinD", "20.870", "0.799", "0.207", "17.544", "0.669"],
        ["PairLIE", "19.510", "0.736", "0.248", "20.357", "0.782"],
        ["SMG", "23.684", "0.826", "0.118", "24.620", "0.867"],
        ["SNR-Net", "24.608", "0.840", "0.151", "21.479", "0.848"],
        ["LLFlow", "24.999", "0.870", "0.117", "26.200", "0.888"],
        ["Retinexformer", "25.153", "0.843", "0.131", "22.794", "0.839"],
        ["Diff-Retinex", "21.981", "0.863", "0.048", "-", "-"],
        ["DiffLL", "26.336", "0.845", "0.217", "28.857", "0.876"],
        ["PyDiff", "27.088", "0.875", "0.111", "27.236", "0.869"],
        ["ReCo-Diff", "27.626", "0.884", "0.090", "29.306", "0.906"],
    ],
}


# ------------------------------------------------------------ main
def clear_cell(cell):
    """remove all inner tables and paragraphs, leave one empty paragraph"""
    for tbl in list(cell.tables):
        tbl._tbl.getparent().remove(tbl._tbl)
    for p in list(cell.paragraphs):
        p._p.getparent().remove(p._p)
    cell.add_paragraph()


def fill_if_empty(cell, text, rtl=False, size=10.5):
    if not cell.text.strip():
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(size)
        if rtl:
            set_rtl(p)
        return True
    return False


def main():
    entries = parse_bib(PROJ / "references.bib")

    sec1 = clean_md((PROJ / "section1-draft.md").read_text(encoding="utf-8").split("## ۱-۱.", 1)[1].join(["## ۱-۱.", ""]))
    all_md = (PROJ / "sections-2-3-4-draft.md").read_text(encoding="utf-8")
    p2 = all_md.split("# (۲)", 1)[1]
    sec2 = clean_md(p2.split("# (۳)", 1)[0])
    sec3 = clean_md(p2.split("# (۳)", 1)[1].split("# (۴)", 1)[0])
    sec4 = clean_md(p2.split("# (۴)", 1)[1])
    md56 = (PROJ / "sections-5-6-draft.md").read_text(encoding="utf-8")
    sec5 = clean_md(md56.split("# (۵)", 1)[1].split("# (۶)", 1)[0])

    items = [sec1, sec2, sec3, sec4, sec5]
    order = build_cite_map(items, entries)
    sec1, sec2, sec3, sec4, sec5 = [apply_cites(s, order) for s in items]

    # map citations inside hardcoded tables too
    for rows in (RISK_ROWS, TIMELINE_ROWS, COMPARISON_ROWS, *TABLES.values()):
        for i, row in enumerate(rows):
            rows[i] = [map_cites(v, order) for v in row]

    doc = Document(str(FORM))
    t1 = doc.tables[1]

    # --- title / keywords: fill only EMPTY cells (row 6 FA title is user-managed) ---
    fill_if_empty(t1.rows[7].cells[1],
                  "A Diffusion-Prior Zero-Shot Method for Low-Light Image Enhancement "
                  "under Imbalanced Twilight Conditions with Application to Machine Vision")
    fill_if_empty(t1.rows[8].cells[1],
                  "بهبود تصویر کم‌نور، مدل‌های انتشار، یادگیری بدون نظارت، گرگ‌ومیش، بینایی ماشین",
                  rtl=True)
    fill_if_empty(t1.rows[9].cells[1],
                  "Low-Light Image Enhancement, Diffusion Models, Zero-Shot Learning, "
                  "Twilight, Machine Vision")

    # --- units / duration (مدت cell located dynamically after its label) ---
    fill_if_empty(t1.rows[11].cells[1], "21")
    cells11 = t1.rows[11].cells
    for i, c in enumerate(cells11):
        if c.text.strip().startswith("مدت اجرا"):
            for j in range(i + 1, len(cells11)):
                if cells11[j]._tc is not c._tc and cells11[j]._tc is not cells11[1]._tc:
                    fill_if_empty(cells11[j], "24 ماه", rtl=True)
                    break
            break

    # --- content cells: clear (incl. inner tables) then refill from sources ---
    clear_cell(t1.rows[23].cells[0])
    fill_cell(t1.rows[23].cells[0], sec1)
    clear_cell(t1.rows[25].cells[0])
    fill_cell(t1.rows[25].cells[0], sec2)
    clear_cell(t1.rows[27].cells[0])
    fill_cell(t1.rows[27].cells[0], sec3)
    clear_cell(t1.rows[29].cells[0])
    fill_cell(t1.rows[29].cells[0], sec4)
    fill_table(t1.rows[29].cells[0], RISK_ROWS)
    add_para(t1.rows[29].cells[0], MILESTONE)

    # novelty cell + comparison table (moved to top-level table 3 by Word re-save)
    t3 = doc.tables[3]
    clear_cell(t3.rows[1].cells[0])
    fill_cell(t3.rows[1].cells[0], sec5)
    fill_table(t3.rows[1].cells[0], COMPARISON_ROWS)

    # refresh the standalone timeline table (table 2, data rows 1..5)
    for ri, row_vals in enumerate(TIMELINE_ROWS[1:], start=1):
        for ci, val in enumerate(row_vals):
            c = doc.tables[2].rows[ri].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci < 2 else WD_ALIGN_PARAGRAPH.JUSTIFY
            set_rtl(p)
            run = p.add_run(val)
            run.font.size = Pt(9)

    # --- references ---
    used = sorted(order.items(), key=lambda kv: kv[1])
    ref_cell = t3.rows[3].cells[0]
    clear_cell(ref_cell)
    for key, num in used:
        line = f"[{num}] {ieee_ref(key, entries)}"
        add_ltr_para(ref_cell, line)

    doc.save(str(OUT))
    return order


if __name__ == "__main__":
    order = main()
    with io.open(r"C:\Users\Alireza\AppData\Local\Temp\opencode\fill_report.txt", "w", encoding="utf-8") as f:
        f.write(f"OK — references used: {len(order)}\n")
        for k, v in sorted(order.items(), key=lambda kv: kv[1]):
            f.write(f"{v}\t{k}\n")
    print("done")
